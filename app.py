from flask import Flask, render_template, request, redirect, url_for, session, send_file, render_template_string, flash
from flask_sqlalchemy import SQLAlchemy
from models import db, Customer, Booking, Tables, MenuItem, OrderItem
from io import BytesIO
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase.ttfonts import TTFont
import google.generativeai as genai
from dotenv import load_dotenv
import json
import os


dotenv_path = r'E:\restaurant_AI\api.env'
load_dotenv(dotenv_path)
api_key = os.getenv("GEMINI_API_KEY")


API_KEY = os.getenv("GEMINI_API_KEY")
print("API_KEY:", API_KEY)

if not API_KEY:
    raise ValueError("GEMINI_API_KEY не знайдено. Додайте у .env")

genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

app = Flask(__name__)
app.secret_key = 'your_secret_key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///new_booking.db'
db.init_app(app)
USERS_FILE = 'users.json'

def load_users():
    if not os.path.exists(USERS_FILE):
        return {}
    with open(USERS_FILE, 'r') as f:
        return json.load(f)

def save_users(users):
    with open(USERS_FILE, 'w') as f:
        json.dump(users, f)
@app.route('/export_booking/<format>')
def export_booking(format):
    bookings = Booking.query.all()

    if format == 'docx':
        doc = Document()
        doc.add_heading('Список бронювань', 0)

        for b in bookings:
            customer = Customer.query.get(b.customer_id)
            order_items = OrderItem.query.filter_by(booking_id=b.id).all()
            dishes = []
            for item in order_items:
                menu_item = MenuItem.query.get(item.menu_item_id)
                if menu_item:
                    dishes.append(menu_item.name)
            dishes_str = ", ".join(dishes) if dishes else "немає"

            doc.add_paragraph(
                f"Ім’я: {customer.name}, Телефон: {customer.phone}, Email: {customer.email}, "
                f"Дата: {b.day}, Час: {b.time}, Гостей: {b.guests}, Стіл: {b.table_id}, Страви: {dishes_str}"
            )

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='booking.docx')

    elif format == 'xlsx':
        wb = Workbook()
        ws = wb.active
        ws.append(['Ім’я', 'Телефон', 'Email', 'Дата', 'Час', 'Гостей', 'Тривалість', 'Стіл', 'Страви'])

        for b in bookings:
            customer = Customer.query.get(b.customer_id)
            order_items = OrderItem.query.filter_by(booking_id=b.id).all()
            dishes = []
            for item in order_items:
                menu_item = MenuItem.query.get(item.menu_item_id)
                if menu_item:
                    dishes.append(menu_item.name)
            dishes_str = ", ".join(dishes) if dishes else "немає"

            ws.append([
                customer.name, customer.phone, customer.email,
                b.day, b.time, b.guests, b.during, b.table_id, dishes_str
            ])

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='booking.xlsx')

    elif format == 'pdf':
        buffer = BytesIO()

        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                leftMargin=2 * cm, rightMargin=2 * cm,
                                topMargin=2 * cm, bottomMargin=2 * cm)

        font_path = os.path.join('static', 'fonts', 'DejaVuSans.ttf')
        pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='MyStyle', fontName='DejaVuSans', fontSize=10, leading=14))
        styles.add(ParagraphStyle(name='MyHeading', fontName='DejaVuSans', fontSize=16, leading=20, alignment=1))

        content = []
        content.append(Paragraph("Список бронювань", styles['MyHeading']))
        content.append(Spacer(1, 12))

        for b in bookings:
            customer = Customer.query.get(b.customer_id)
            order_items = OrderItem.query.filter_by(booking_id=b.id).all()
            dishes = []
            for item in order_items:
                menu_item = MenuItem.query.get(item.menu_item_id)
                if menu_item:
                    dishes.append(menu_item.name)
            dishes_str = ", ".join(dishes) if dishes else "немає"

            text = f"""
            <b>Ім’я:</b> {customer.name}, <b>Телефон:</b> {customer.phone}, <b>Email:</b> {customer.email}<br/>
            <b>Дата:</b> {b.day}, <b>Час:</b> {b.time}, <b>Гостей:</b> {b.guests}, <b>Тривалість:</b> {b.during}, <b>Стіл:</b> {b.table_id}<br/>
            <b>Страви:</b> {dishes_str}
            """
            content.append(Paragraph(text, styles['MyStyle']))
            content.append(Spacer(1, 8))

        doc.build(content)
        buffer.seek(0)

        return send_file(buffer, as_attachment=True, download_name="booking.pdf")



@app.route('/')
def index():
    return render_template('main.html')

@app.route('/menu')
def menu():
    menu_items = [
        {
            "id": 1,
            "title": "Стейк зі свинини на грилі зі смаженою картоплею",
            "price": 289,
            "image_url": "/static/uploads/3l0GvTy9P8lMAjho8bdBK4f37C9qqdyJg18zX5fX.webp"
        },
        {
            "id": 2,
            "title": "Млинці з м'ясом та томатним соусом",
            "price": 188,
            "image_url": "/static/uploads/7fYwur0eTaKgc51iLkptgxemwBML8qk82KlExi7V.webp"
        },
        {
            "id": 3,
            "title": "Паста Карбонара",
            "price": 176,
            "image_url": "static/uploads/Lhm6HJbNA8eDlzMHQuJT2Yit8gcEo7FLxi1YW0o6.webp"
        },
        {
            "id": 4,
            "title": "Стейк лосося з кус-кусом в соусі Песто",
            "price": 289,
            "image_url": "static/uploads/URRmjVgmx3lhwvcZFNRNDHk2JCxGlcyiSVgI8JLz.webp"
        },
        {
            "id": 5,
            "title": "Паста у томатному соусі",
            "price": 210,
            "image_url": "static/uploads/xKwzERFOQOAfILBh0lvpLaW1mtl6fxEZWXcfQUq2.webp"
        },
        {
            "id": 6,
            "title": "Coca-cola",
            "price": 54,
            "image_url": "static/uploads/FvHplte-bbCyvnk-ZJjreGh.webp"
        },
        {
            "id": 7,
            "title": "Вода Bonaqua",
            "price": 44,
            "image_url": "static/uploads/pCuMoDz-MvqddzF-fzodjjw.webp"
        },
        {
            "id": 8,
            "title": "Салат Грецький",
            "price": 144,
            "image_url": "static/uploads/NAwZlcM-TYIbKyk-gsKwehF.webp"
        }

    ]
    return render_template('menu.html', menu_items=menu_items)

@app.route('/add_to_cart', methods=['POST'])
def add_to_cart():
    item_id = request.form.get('item_id')
    title = request.form.get('title')
    price = float(request.form.get('price'))

    cart = session.get('cart', [])

    for item in cart:
        if item['title'] == title:
            item['quantity'] += 1  
            break
    else:
        cart.append({'item_id': int(item_id), 'title': title, 'price': price, 'quantity': 1})

    session['cart'] = cart
    session.modified = True


    return redirect(url_for('menu'))


@app.route('/cart')
def cart():
    cart_items = session.get('cart', [])
    return render_template('cart.html', cart=cart_items)

@app.route('/remove_from_cart/<int:item_index>', methods=['POST'])
def remove_from_cart(item_index):
    cart = session.get('cart', [])
    if 0 <= item_index < len(cart):
        cart.pop(item_index)
    session['cart'] = cart
    session.modified = True
    return redirect(url_for('cart'))

@app.route('/authorize', methods=['GET', 'POST'])
def authorize():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        users = load_users()

        if username in users and users[username]['password'] == password:
            session['user'] = username
            session['role'] = users[username]['role']
            if session['role'] == 'admin':
                return redirect(url_for('show'))
            else:
                return redirect(url_for('index'))
    return render_template('authorize.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        users = load_users()

        if username in users:
            return "Користувач вже існує!"
        
        users[username] = {
            "password": password,
            "role": "user"
        }

        save_users(users)
        return redirect(url_for('authorize'))
    return render_template('register.html')

@app.route('/home')
def home():
    if 'user' not in session:
        return redirect(url_for('authorize'))
    return render_template('main.html', user=session['user'])

@app.route('/booking', methods=['GET', 'POST'])
def booking():
    if request.method == 'POST':
        session['day'] = request.form.get('day')
        session['time'] = request.form.get('time')
        session['guests'] = request.form.get('guests')
        session['during'] = request.form.get('during')
        session['table_id'] = request.form.get('table_id')
        return redirect(url_for('contacts'))

    cart_items = session.get('cart', [])
    tables = Tables.query.all()
    return render_template('booking.html', tables=tables, cart=cart_items)


@app.route('/contacts')
def contacts():
    return render_template('contacts.html')

@app.route('/add_contacts', methods=['POST'])
def add_contacts():
    required_fields = ['day', 'time', 'guests', 'during', 'table_id']
    for field in required_fields:
        if not session.get(field):
            flash(f"Відсутні дані для {field}. Будь ласка, заповніть форму знову.")
            return redirect(url_for('booking'))

    try:
        guests = int(session.get('guests'))
        table_id = int(session.get('table_id'))
    except (TypeError, ValueError):
        flash("Некоректні дані для гостей або столу.")
        return redirect(url_for('booking'))

    customer = Customer(
        name=request.form.get('name'),
        phone=request.form.get('tel'),
        email=request.form.get('email')
    )
    db.session.add(customer)
    db.session.commit()

    booking = Booking(
        day=session.get('day'),
        time=session.get('time'),
        guests=guests,
        during=session.get('during'),
        table_id=table_id,
        customer_id=customer.id
    )
    db.session.add(booking)
    db.session.commit()
    

    cart = session.get('cart', [])
    for item in cart:
        order_item = OrderItem(
            booking_id=booking.id,
            menu_item_id=item['item_id'],
            quantity=item['quantity']
        )
        db.session.add(order_item)
    db.session.commit()
    print("OrderItems in DB for booking", booking.id, OrderItem.query.filter_by(booking_id=booking.id).all())


    customer_name = customer.name
    customer_email = customer.email
    table_number = Tables.query.get(table_id).table_number if Tables.query.get(table_id) else "Невідомий"

    session.clear()

    return redirect(url_for('confirmation', 
                            name=customer_name, 
                            email=customer_email, 
                            table_number=table_number))


@app.route('/confirmation')
def confirmation():
    name = request.args.get('name')
    email = request.args.get('email')
    table_number = request.args.get('table_number')
    return render_template('booking_result.html', name=name, email=email, table_number=table_number)


@app.route('/show')
def show():
    if 'user' not in session or session.get('role') != 'admin':
        flash("Будь ласка, увійдіть, щоб переглянути резервації.")
        return redirect(url_for('authorize'))

    query = Booking.query

    name = request.args.get('name')
    day = request.args.get('day')
    table_id = request.args.get('table_id')

    if name:
        query = query.join(Customer).filter(Customer.name.ilike(f'%{name}%'))
    if day:
        query = query.filter(Booking.day == day)
    if table_id:
        query = query.filter(Booking.table_id == table_id)

    customers = {c.id: c for c in Customer.query.all()}
    order_items = OrderItem.query.all()
    menu_items = {m.id: m for m in MenuItem.query.all()}

    orders_by_booking = {}
    for item in order_items:
        menu_item = menu_items.get(item.menu_item_id)
        if menu_item:
            orders_by_booking.setdefault(str(item.booking_id), []).append({
                'title': menu_item.name,
                'quantity': item.quantity,
                'price':  menu_item.price
            })

    bookings = query.all()
    tables = Tables.query.all()

    return render_template('show.html', reservations=bookings, tables=tables, orders=orders_by_booking,  customers=customers,)


@app.route('/summary')
def summary():
    if 'user' not in session or session.get('role') != 'admin':
        flash("Будь ласка, увійдіть, щоб переглянути резервації.")
        return redirect(url_for('authorize'))

    bookings = Booking.query.all()
    total_guests = sum(b.guests for b in bookings)
    total_bookings = len(bookings)

    orders = OrderItem.query.all()
    dish_amount = {}
    for o in orders:
        item = MenuItem.query.get(o.menu_item_id)
        if item:
            dish_amount[item.name] = dish_amount.get(item.name, 0) + o.quantity    
    top_dishes = sorted(dish_amount.items(), key=lambda x: -x[1])[:3]

    text = (
        f"Було {total_bookings} бронювань на {total_guests} гостей. "
        f"Топ страви: " + ", ".join(f"{name} ({count})" for name, count in top_dishes)
    )

    summary_text = generate_summary(text)

    return render_template("summary_for_admin.html", summary=summary_text)
        

def generate_summary(prompt):
    try:
        question = f"Сформуй коротке підсумкове резюме для адміністратора: {prompt}"
        response = model.generate_content(question)

        if response and response.text:
            return response.text
        else:
            print(f"Gemini API повернув порожню відповідь для запиту: {prompt}")
            return "Не вдалося згенерувати підсумок. Спробуйте пізніше або змініть запит."    
    except Exception as error:
        print(f"Помилка при запиті Gemini API: {error}")
        return f"Помилка генерації підсумку: {error}. Оригінальний текст: {prompt}"

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))

@app.before_request
def init_menu_items():
    if not MenuItem.query.first(): 
        db.session.add_all([
            MenuItem(id=1, name="Стейк зі свинини на грилі зі смаженою картоплею", price=289),
            MenuItem(id=2, name="Млинці з м'ясом та томатним соусом", price=188),
            MenuItem(id=3, name="Паста Карбонара", price=176),
            MenuItem(id=4, name="Стейк лосося з кус-кусом в соусі Песто", price=289),
            MenuItem(id=5, name="Паста у томатному соусі", price=210),
            MenuItem(id=6, name="Coca-cola", price=54),
            MenuItem(id=7, name="Вода Bonaqua", price=44),
        ])
        db.session.commit()


@app.before_request
def init_tables():
    if not Tables.query.first():
        db.session.add_all([
            Tables(table_number="1", seats=2),
            Tables(table_number="2", seats=4),
            Tables(table_number="3", seats=6),
            Tables(table_number="4", seats=8),
        ])
        db.session.commit()
        print("Столи успішно додані!")

@app.before_request
def create_views_and_triggers():
    with open('static/create_view_and_trigers.sql', 'r', encoding='utf-8') as f:
        sql_text = f.read().split(';')
    for comand in sql_text:
        comand = comand.strip()
    if comand:  
        db.session.execute(comand)
    db.session.commit()

if __name__ == "__main__":
    with app.app_context():
        db.create_all()
        init_tables()
        init_menu_items()
        create_views_and_triggers()
    app.run(debug=True)
